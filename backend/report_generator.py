from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import sqlite3
import json
import os


def generate_mayor_report():
    """Генерация отчета для Главы в формате .docx"""
    try:
        # Подключение к БД
        db_path = '../data/municipal_monitoring.db'
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()

        # Получаем статистику за 24 часа
        cursor.execute('''
            SELECT 
                COUNT(*) as total,
                SUM(CASE WHEN priority >= 2 THEN 1 ELSE 0 END) as critical,
                category, COUNT(*) as count
            FROM problems 
            WHERE created_at > datetime('now', '-24 hours')
            GROUP BY category
            ORDER BY count DESC
        ''')

        categories = cursor.fetchall()

        # Получаем критические проблемы
        cursor.execute('''
            SELECT text, category, location, priority, created_at
            FROM problems 
            WHERE priority >= 2 AND created_at > datetime('now', '-24 hours')
            ORDER BY priority DESC
            LIMIT 10
        ''')

        critical_problems = cursor.fetchall()
        conn.close()

        # Создаем документ
        doc = Document()

        # Заголовок
        title = doc.add_heading('ЕЖЕДНЕВНАЯ СВОДКА', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Подзаголовок
        subtitle = doc.add_heading(f'AI-помощник Главы Екатеринбурга\n{datetime.now().strftime("%d.%m.%Y %H:%M")}', 1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Раздел 1: Статистика
        doc.add_heading('📊 Статистика за 24 часа', level=2)

        if categories:
            total = sum([c[2] for c in categories])
            critical = sum([c[1] for c in categories])

            stats_table = doc.add_table(rows=1, cols=3)
            stats_table.style = 'Light Grid Accent 1'

            hdr_cells = stats_table.rows[0].cells
            hdr_cells[0].text = 'Всего проблем'
            hdr_cells[1].text = 'Критических'
            hdr_cells[2].text = 'Категорий'

            row_cells = stats_table.add_row().cells
            row_cells[0].text = str(total)
            row_cells[1].text = str(critical)
            row_cells[2].text = str(len(categories))

        # Раздел 2: Критические проблемы
        doc.add_heading('🚨 Критические проблемы', level=2)

        if critical_problems:
            crit_table = doc.add_table(rows=1, cols=4)
            crit_table.style = 'Light Grid Accent 1'

            hdr_cells = crit_table.rows[0].cells
            hdr_cells[0].text = 'Категория'
            hdr_cells[1].text = 'Место'
            hdr_cells[2].text = 'Описание'
            hdr_cells[3].text = 'Приоритет'

            for problem in critical_problems:
                row_cells = crit_table.add_row().cells
                row_cells[0].text = problem[1]
                row_cells[1].text = problem[2]
                row_cells[2].text = problem[0][:100] + '...'
                row_cells[3].text = str(problem[3])

        # Раздел 3: Топ категорий
        doc.add_heading('🏆 Топ категорий проблем', level=2)

        if categories:
            top_table = doc.add_table(rows=1, cols=2)
            top_table.style = 'Light Grid Accent 1'

            hdr_cells = top_table.rows[0].cells
            hdr_cells[0].text = 'Категория'
            hdr_cells[1].text = 'Количество'

            for category in categories[:5]:  # Только топ-5
                row_cells = top_table.add_row().cells
                row_cells[0].text = category[0]
                row_cells[1].text = str(category[2])

        # Раздел 4: Рекомендации
        doc.add_heading('🎯 Рекомендации', level=2)
        recommendations = [
            "1. Немедленно отреагировать на критические проблемы (приоритет 2+)",
            "2. Усилить мониторинг наиболее частых категорий",
            "3. Координировать работу служб по кластерам проблем",
            "4. Проверить выполнение предыдущих поручений",
            "5. Подготовить пресс-релиз о принимаемых мерах"
        ]

        for rec in recommendations:
            doc.add_paragraph(rec)

        # Сохраняем документ
        reports_dir = '../reports'
        os.makedirs(reports_dir, exist_ok=True)

        filename = f"mayor_report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        filepath = os.path.join(reports_dir, filename)
        doc.save(filepath)

        print(f"✅ Отчет сгенерирован: {filepath}")
        return filepath

    except Exception as e:
        print(f"❌ Ошибка генерации отчета: {e}")
        return None


if __name__ == "__main__":
    generate_mayor_report()